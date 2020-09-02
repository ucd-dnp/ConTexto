# Clase escritor

class Escritor():
    def __init__(self, ubicacion_archivo, texto):
        """
        Constructor por defecto de la clase Escritor. Esta clase se encarga de guardar \
            texto en archivos de distintos tipos como Word, PDF, CSV, TXT, RTF e \
            imágenes
        :param ubicacion_archivo: (string). Ruta del archivo que será guardado \
            con el texto deseado
        :param texto: (string). Texto que se desea guardar en un archivo
        """
        self.establecer_ubicacion(ubicacion_archivo)
        self.establecer_texto(texto)

    def establecer_ubicacion(self, ubicacion_archivo):
        """ Define la ruta del archivo con el texto que se desea guardar 

        :param ubicacion_archivo: (string). Ruta del archivo que será guardado \
            con el texto deseado 
        """
        self.ubicacion_archivo = ubicacion_archivo

    def establecer_texto(self, texto):
        """ Define el texto que será guardado en un archivo

        :param texto: (string). Texto que será guardado en un archivo 
        """
        self.texto = texto

    def escribir_txt(self):
        """ Especifica que el texto deseado se guardará en un archivo con extensión '.txt'
        """
        if isinstance(self.texto, list):
            self.texto = '\n\n|**|\n\n'.join(self.texto)
        # with open(self.ubicacion_archivo, 'w') as fp:
        with open(self.ubicacion_archivo, 'w', encoding="utf-8") as fp:
            fp.write(self.texto)

    def escribir_word(self):
        """ Especifica que el texto deseado se guardará en un archivo con extensión '.docx'
        """
        from docx import Document
        documento = Document()
        if isinstance(self.texto, list):
            for i, page in enumerate(self.texto):
                documento.add_paragraph(page)
                if i < len(self.texto) - 1:
                    documento.add_page_break()
        else:
            documento.add_paragraph(self.texto)
        documento.save(self.ubicacion_archivo)

    def escribir_pdf(self):
        """ Especifica que el texto deseado se guardará en un archivo con extensión '.pdf'
        """
        import PyPDF2
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO
        from textwrap import wrap

        def escribir_pagina(texto):

            texto = texto.split('\n')
            temp = BytesIO()
            can = canvas.Canvas(temp, pagesize=letter)
            t = can.beginText()
            t.setFont('Helvetica-Bold', 7)
            t.setCharSpace(0)
            t.setTextOrigin(50, 700)
            for linea in texto:
                sublineas = wrap(linea, 150)
                if len(sublineas) > 0:
                    for sublinea in sublineas:
                        t.textLine(sublinea)
                else:
                    t.textLine('')
            can.drawText(t)
            # can.drawString(5, 550, text)
            can.save()
            temp.seek(0)
            lector = PyPDF2.PdfFileReader(temp)
            return lector
        salida = PyPDF2.PdfFileWriter()
        if isinstance(self.texto, list):
            for pag in self.texto:
                lector = escribir_pagina(pag)
                salida.addPage(lector.getPage(0))
        else:
            lector = escribir_pagina(self.texto)
            salida.addPage(lector.getPage(0))
        with open(self.ubicacion_archivo, 'wb') as fp:
            salida.write(fp)

    def texto_a_archivo(self, tipo='inferir'):
        """ Especifica el tipo de archivo en el que se quiere guardar el texto

        :param tipo: (string) {'inferir', 'txt', 'csv', 'pdf', 'doc', 'docx'}. Valor por defecto: 'inferir'. \
            Define el tipo del archivo en el que se desea guardar el texto
        """
        if tipo == 'inferir':
            tipo = self.ubicacion_archivo.split('.')[-1]
        if tipo in ['txt', 'csv']:
            self.escribir_txt()
        elif tipo == 'pdf':
            self.escribir_pdf()
        elif tipo in ['doc', 'docx']:
            self.escribir_word()
        else:
            print('Formato desconocido. Se escribirá en un formato plano (.txt).')
            nueva_ruta = ''.join(self.ubicacion_archivo.split(
                '.')[:-1]) + '_{}.txt'.format(tipo)
            self.establecer_ubicacion(nueva_ruta)
            self.escribir_txt()


# Función que encapsula el proceso de escritura de archivos de texto

def escribir_texto(ubicacion_archivo, texto, tipo='inferir'):
    """ Función que guarda texto en un archivo específico. Permite escoger la \
        ruta del archivo, su tipo y el texto a ser guardado dentro de este archivo

    :param ubicacion_archivo: (string). Ruta del archivo que será guardado \
        con el texto deseado
    :param texto: (string). El texto que se desea guardar en un archivo
    :param tipo: (string) {'inferir', 'txt', 'csv', 'pdf', 'doc', 'docx'}. Valor por defecto: 'inferir'. \
         Define el tipo del archivo en el que se desea guardar el texto
    """
    es = Escritor(ubicacion_archivo, texto)
    es.texto_a_archivo(tipo)
